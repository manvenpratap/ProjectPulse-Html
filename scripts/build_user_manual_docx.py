"""
build_user_manual_docx.py
=========================
Generates the ProjectPulse Product Capabilities & User Manual as a single
production-quality DOCX with:
  • Branded cover page
  • Table of Contents
  • 10 fully documented chapters (one per module/view)
  • Actual application screenshots embedded as figures
  • Data tables, callout boxes, step-by-step procedures
  • Consistent typography and professional layout
"""

import os, datetime, glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ──────────────────────────────────────────────────────────────────
BASE_DIR   = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
SS_DIR     = os.path.join(BASE_DIR, "docs", "screenshots")
OUT_PATH   = os.path.join(BASE_DIR, "docs", "ProjectPulse_Product_Capabilities_Manual.docx")
LOGO_PATH  = None   # set if you have a logo PNG

# ── Brand Palette ──────────────────────────────────────────────────────────
CLR_NAVY    = RGBColor(0x0D, 0x1B, 0x2A)   # Deep navy heading
CLR_TEAL    = RGBColor(0x00, 0xC2, 0xA8)   # Accent teal
CLR_ACCENT  = RGBColor(0x1E, 0x9E, 0xD5)   # Sky blue
CLR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
CLR_LIGHT   = RGBColor(0xF4, 0xF7, 0xFA)   # Light grey bg
CLR_DARK    = RGBColor(0x1A, 0x1A, 0x2E)
CLR_WARN    = RGBColor(0xE6, 0x7E, 0x22)   # Amber warning
CLR_DANGER  = RGBColor(0xC0, 0x39, 0x2B)   # Red
CLR_SUCCESS = RGBColor(0x27, 0xAE, 0x60)   # Green
CLR_TEXT    = RGBColor(0x1C, 0x1C, 0x1E)   # Primary text
CLR_MUTED   = RGBColor(0x6C, 0x75, 0x7D)   # Secondary text
CLR_BORDER  = RGBColor(0xD1, 0xD5, 0xDB)   # Table borders

# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set cell background shading via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), kwargs.get(edge, {}).get("val", "single"))
        tag.set(qn("w:sz"), str(kwargs.get(edge, {}).get("sz", 4)))
        tag.set(qn("w:color"), kwargs.get(edge, {}).get("color", "D1D5DB"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc, color="D1D5DB"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p


def make_heading(doc, text, level=1, color=None, space_before=18, space_after=6):
    """Add a styled heading paragraph."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(space_after)
    run = h.runs[0] if h.runs else h.add_run(text)
    if color:
        run.font.color.rgb = color
    return h


def make_body(doc, text, space_before=0, space_after=6, color=None, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if bold:  run.bold = True
    if italic: run.italic = True
    return p


def make_bullet(doc, text, level=0, space_after=3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def make_numbered(doc, text, level=0, space_after=3):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_screenshot(doc, filename, caption, width_inches=6.3):
    """Embed a screenshot from SS_DIR with a styled caption."""
    # Try exact name, then glob
    candidates = [
        os.path.join(SS_DIR, filename),
        os.path.join(SS_DIR, filename + ".png"),
    ]
    path = None
    for c in candidates:
        if os.path.exists(c):
            path = c
            break
    if not path:
        # Try glob prefix match
        matches = glob.glob(os.path.join(SS_DIR, f"{filename}*"))
        if matches:
            path = sorted(matches)[0]

    if path and os.path.exists(path):
        p_img = doc.add_paragraph()
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after  = Pt(0)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(path, width=Inches(width_inches))
    else:
        # Placeholder
        p_img = doc.add_paragraph(f"[Screenshot: {filename}]")
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.runs[0].font.color.rgb = CLR_MUTED
        p_img.runs[0].font.italic = True

    # Caption
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(14)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = CLR_MUTED
    return p_img


def add_callout(doc, text, style="note"):
    """Add a styled callout / note box."""
    configs = {
        "note":    ("ℹ  NOTE",    "0D4F8B", "E8F4FD", "1565C0"),
        "tip":     ("💡 TIP",     "145A32", "E9F7EF", "1E8449"),
        "warning": ("⚠  WARNING", "7D4E00", "FEF9E7", "C67C00"),
        "caution": ("⛔ CAUTION", "6E2222", "FDEDEC", "C0392B"),
    }
    label, label_col, bg_col, border_col = configs.get(style, configs["note"])

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_col)
    # Left border accent
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, col, sz in [("left", border_col, 18), ("top", "FFFFFF", 4), ("right", "FFFFFF", 4), ("bottom", "FFFFFF", 4)]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(sz))
        tag.set(qn("w:color"), col)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.1)
    lbl_run = p.add_run(label + "  ")
    lbl_run.bold = True
    lbl_run.font.size = Pt(9)
    lbl_run.font.color.rgb = RGBColor.from_string(label_col)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(6)
    p2.paragraph_format.left_indent  = Inches(0.1)
    body_run = p2.add_run(text)
    body_run.font.size = Pt(10)
    doc.add_paragraph()
    return tbl


def add_data_table(doc, headers, rows, col_widths=None, header_bg="0D1B2A", header_fg="FFFFFF"):
    """Add a styled data table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(header_fg)

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        bg = "F8F9FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            set_cell_bg(cells[c_idx], bg)
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            if p.runs:
                p.runs[0].font.size = Pt(9.5)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


def add_page_break(doc):
    doc.add_page_break()


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


# ── Cover Page ─────────────────────────────────────────────────────────────

def build_cover(doc):
    doc.add_paragraph("\n\n\n")

    # Product name
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("ProjectPulse")
    run.bold = True
    run.font.size = Pt(44)
    run.font.color.rgb = CLR_NAVY

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Product Capabilities & User Manual")
    run.font.size = Pt(20)
    run.font.color.rgb = CLR_TEAL

    doc.add_paragraph("\n")

    # Horizontal rule
    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hr.add_run("─" * 60)
    run.font.color.rgb = CLR_TEAL
    run.font.size = Pt(12)

    doc.add_paragraph("\n")

    # Description
    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc_p.add_run(
        "The definitive functional reference guide covering every module, feature,\n"
        "workflow, and configuration option in the ProjectPulse workstation."
    )
    run.font.size = Pt(13)
    run.font.color.rgb = CLR_MUTED

    doc.add_paragraph("\n\n")

    # Meta block
    meta_items = [
        ("Document Type", "Product Capabilities & User Manual"),
        ("Version",       "v2.1.0"),
        ("Audience",      "Project Leads · Executives · Developers · QA Engineers"),
        ("Date",          datetime.date.today().strftime("%d %B %Y")),
        ("Confidentiality", "Internal Use — Commercial in Confidence"),
    ]
    tbl = doc.add_table(rows=len(meta_items), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta_items):
        c0, c1 = tbl.rows[i].cells
        set_cell_bg(c0, "0D1B2A")
        set_cell_bg(c1, "F4F7FA")
        c0.width = Inches(2.0)
        c1.width = Inches(3.5)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(label)
        r0.bold = True; r0.font.size = Pt(9.5)
        r0.font.color.rgb = CLR_WHITE
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = CLR_TEXT

    add_page_break(doc)


# ── Table of Contents (manual) ─────────────────────────────────────────────

def build_toc(doc):
    make_heading(doc, "Table of Contents", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")

    chapters = [
        ("1",   "Executive Overview Dashboard",                 "5"),
        ("2",   "Live Project Insights & Analytics",            "10"),
        ("3",   "Hierarchical Delivery Matrix",                 "15"),
        ("4",   "Gantt Timeline & Dependency Management",       "22"),
        ("5",   "Intelligent Weekly Scheduler & Conflict Resolver", "28"),
        ("6",   "Unified RAID Register",                        "35"),
        ("7",   "Team Capacity Hub & Role Management",          "41"),
        ("8",   "Defect Tracker & Bug Lifecycle",               "46"),
        ("9",   "Reports & Stakeholder Board Packs",            "52"),
        ("10",  "Audit Log & Activity Streams",                 "56"),
        ("11",  "Schedule Baselines & Snapshot History",        "60"),
        ("12",  "System Configuration Reference",               "65"),
        ("A",   "Appendix A — Keyboard Shortcuts",              "70"),
        ("B",   "Appendix B — Data Schema Reference",           "72"),
    ]

    tbl = doc.add_table(rows=len(chapters), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (num, title, pg) in enumerate(chapters):
        c0, c1, c2 = tbl.rows[i].cells
        c0.width = Inches(0.4)
        c1.width = Inches(5.4)
        c2.width = Inches(0.6)
        for cell in (c0, c1, c2):
            bg = "F4F7FA" if i % 2 == 0 else "FFFFFF"
            set_cell_bg(cell, bg)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(num)
        r0.bold = True; r0.font.size = Pt(10)
        r0.font.color.rgb = CLR_TEAL
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(title)
        r1.font.size = Pt(10.5)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(pg)
        r2.font.size = Pt(10)
        r2.font.color.rgb = CLR_MUTED

    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER CONTENT
# ══════════════════════════════════════════════════════════════════════════════

def ch_overview(doc):
    make_heading(doc, "1  Executive Overview Dashboard", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The Executive Overview Dashboard is the command centre of ProjectPulse. "
        "It surfaces the three core health signals of the project — Health Index, "
        "Delivery Confidence, and Active RAID Count — in a single, always-visible "
        "pane. Every metric on this screen is computed in real time from the live "
        "state of tasks, resources, risks, and defects without requiring a page "
        "refresh or manual recalculation.",
        space_after=8)

    add_screenshot(doc, "01_overview_dashboard",
                   "Figure 1.1 — Executive Overview Dashboard with KPI cards, Health Gauge, and Predictive Sandbox")

    # ── 1.1 KPI Cards
    make_heading(doc, "1.1  Key Performance Indicator Cards", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Three large KPI cards occupy the top ribbon of the dashboard. Each card "
        "updates automatically whenever underlying data changes.", space_after=6)

    add_data_table(doc,
        ["Metric", "Description", "Calculation Logic", "Target Threshold"],
        [
            ["Health Index",
             "Composite score reflecting overall project vitality.",
             "100 − (Overdue/Active × 50) − (Blocked/Active × 30) − (On Hold/Active × 20)",
             "≥ 80% — Green"],
            ["Delivery Confidence",
             "Probability of delivering on schedule based on team velocity vs. remaining backlog.",
             "Matches remaining task count against 4-week rolling average velocity. Deficit degrades confidence to a floor of 30%.",
             "≥ 85% — Green"],
            ["Active RAID Elements",
             "Total open Risk, Assumption, Issue, and Dependency items.",
             "COUNT(raids WHERE status ∈ {Identified, Active})",
             "0 — Goal"],
        ],
        col_widths=[1.4, 1.9, 2.3, 1.2],
    )

    add_callout(doc,
        "The Health Index excludes Completed and Cancelled tasks from its 'Active' denominator. "
        "Only tasks that are genuinely in-flight and at risk contribute to health degradation.",
        style="note")

    # ── 1.2 Health Gauge
    make_heading(doc, "1.2  Dynamic Health Gauge", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Below the KPI cards sits the radial SVG Health Gauge — a premium animated "
        "dial that provides an at-a-glance summary of the Health Index. The gauge "
        "uses a dynamic arc drawn in real time by the browser's SVG engine, with a "
        "glow effect that intensifies as health improves.", space_after=6)

    make_bullet(doc, "Green glow (80–100%): Project is healthy and on track.")
    make_bullet(doc, "Amber arc (50–79%): Warning — one or more risk factors require attention.")
    make_bullet(doc, "Red arc (<50%): Critical — immediate management intervention required.")
    make_bullet(doc, "Hover interaction: Reveals a breakdown tooltip listing each drag factor and its percentage contribution (e.g., 'Overdue Tasks: −8%', 'Defects: −12%').")
    doc.add_paragraph()

    # ── 1.3 Sandbox
    make_heading(doc, "1.3  What-If Predictive Sandbox", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The Predictive Sandbox is a non-destructive simulation environment that "
        "lets project managers explore 'what-if' scenarios without altering the live "
        "project state. The sandbox operates on an in-memory clone of the master "
        "schedule (P.sandboxTasks), leaving P.tasks completely unchanged.", space_after=6)

    add_data_table(doc,
        ["Control", "Action", "Effect on Simulation"],
        [
            ["Toggle Sandbox", "Activates/deactivates sandbox mode.", "Switches state context from master to sandbox clone."],
            ["Capacity Slider", "Reduce team availability by a percentage (e.g. −20%).", "Extends forecast due dates proportionally across assigned tasks."],
            ["Scope Multiplier", "Increase task complexity by a factor (e.g. 1.2×).", "Inflates estimated effort and pushes projected delivery dates outward."],
            ["Commit Changes", "Applies sandbox mutations to the live project.", "Overwrites master schedule with simulated dates. Irreversible without a baseline restore."],
            ["Discard", "Rolls back all sandbox mutations.", "Restores the live schedule to its pre-simulation state instantly."],
        ],
        col_widths=[1.5, 2.2, 2.7],
    )

    add_callout(doc,
        "Always capture a Schedule Baseline (Chapter 11) before committing sandbox changes. "
        "This ensures you can restore the original plan if the simulation produces undesirable outcomes.",
        style="warning")

    add_page_break(doc)


def ch_insights(doc):
    make_heading(doc, "2  Live Project Insights & Analytics", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The Insights tab is the analytical engine of ProjectPulse. It transforms "
        "raw task logs and effort records into actionable intelligence: Earned Value "
        "Management (EVM) metrics, burn-up projections, velocity sparklines, and "
        "30-day activity heatmaps. All computations are powered by the "
        "buildDashCache() function — an O(n) single-pass aggregator that processes "
        "the full task and log dataset once to produce O(1) lookup structures "
        "that power every widget.", space_after=8)

    add_screenshot(doc, "02_insights_analytics",
                   "Figure 2.1 — Live Insights: EVM metrics, Burn-up Chart, and Velocity Dashboard")

    make_heading(doc, "2.1  Performance Metric Caching (buildDashCache)", level=2, color=CLR_ACCENT)
    make_body(doc,
        "To ensure smooth animations and sub-millisecond render times, ProjectPulse "
        "caches all computed analytics in a single P.cache object. The cache is "
        "hydrated every time task data or resource allocations change, and before "
        "every render cycle.", space_after=6)
    make_bullet(doc, "Single-pass O(n) aggregation — iterates the task list exactly once per cache build.")
    make_bullet(doc, "Calculates time-series arrays for cumulative completion tracking.")
    make_bullet(doc, "Pre-computes dependency lookup maps to eliminate redundant nested searches.")
    make_bullet(doc, "Generates sparkline data points for all health metric history charts.")
    doc.add_paragraph()

    make_heading(doc, "2.2  Earned Value Management (EVM)", level=2, color=CLR_ACCENT)
    make_body(doc,
        "ProjectPulse implements the full PMI-standard EVM framework. Metrics are "
        "displayed as live cards updated on every cache refresh.", space_after=6)

    add_data_table(doc,
        ["Metric", "Formula", "Interpretation"],
        [
            ["Planned Value (PV)", "Total Budget × (Time Elapsed / Total Timeline)", "Budget that should have been consumed by today's date per the original plan."],
            ["Earned Value (EV)", "Total Budget × % Completion", "Value of work actually delivered relative to the plan."],
            ["Actual Cost (AC)", "Σ(logged hours) × avg. blended resource rate", "Real resources consumed to date."],
            ["Schedule Variance (SV)", "EV − PV", "Positive = ahead of schedule. Negative = behind schedule."],
            ["Schedule Performance Index (SPI)", "EV ÷ PV", "SPI < 1.0 indicates schedule slippage. SPI > 1.0 indicates ahead-of-schedule delivery."],
            ["Cost Variance (CV)", "EV − AC", "Positive = under budget. Negative = over budget."],
        ],
        col_widths=[1.8, 2.4, 2.2],
    )

    make_heading(doc, "2.3  Burn-up & Burn-down Projections", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The central chart in the Insights view is an interactive burn-up chart that "
        "provides a visual projection of project completion.", space_after=6)
    make_bullet(doc, "Scope Line (Blue): Total planned complexity points over time. Rises when scope is added.")
    make_bullet(doc, "Completion Line (Green): Accumulated complexity points of all tasks with status 'Completed'. Rising steeply indicates a healthy delivery pace.")
    make_bullet(doc, "Dynamic Projection (Dashed Amber): Linear extrapolation based on the rolling 3-week velocity, estimating the exact calendar week of project completion.")
    doc.add_paragraph()

    make_heading(doc, "2.4  30-Day Activity Heatmap", level=2, color=CLR_ACCENT)
    make_body(doc,
        "A calendar heatmap shows contribution density across the past 30 calendar "
        "days. Each cell represents a single day; colour intensity reflects the "
        "number of task state changes logged by the team on that day. This view is "
        "powered by the retrospective activity date resolution system introduced in "
        "the June 2026 release — log entries reference actCompletionDate rather than "
        "insertion timestamps, giving an accurate picture of when work was actually "
        "completed rather than when it was logged.", space_after=6)

    add_callout(doc,
        "Velocity is calculated as the 3-week rolling average of completed complexity points per week. "
        "Weeks where no tasks were completed are excluded from the rolling window to avoid artificially "
        "suppressing the velocity figure during holiday periods.",
        style="tip")

    add_page_break(doc)


def ch_delivery(doc):
    make_heading(doc, "3  Hierarchical Delivery Matrix", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The Delivery Matrix is the operational backbone of ProjectPulse — "
        "a high-density, spreadsheet-style task grid that supports multi-level "
        "parent-child hierarchies, inline editing, dependency mapping, effort "
        "estimation, and real-time status tracking. Every deliverable, epic, "
        "task, and subtask in the project lives in this view.", space_after=8)

    add_screenshot(doc, "02_delivery_matrix",
                   "Figure 3.1 — Hierarchical Delivery Matrix showing parent tasks, child subtasks, status pills, and inline controls")

    make_heading(doc, "3.1  Parent-Child Task Hierarchies", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Tasks can be structured into multi-level trees. Parent tasks (Epics or "
        "Deliverables) aggregate metrics from all child tasks:", space_after=6)
    make_bullet(doc, "Parent Tasks: Display summary progress as the weighted average of child task progress.")
    make_bullet(doc, "Child Tasks / Subtasks: Individual action items with their own assignees, dates, effort, and complexity.")
    make_bullet(doc, "Expand/Collapse: Click the chevron icon on any parent row to toggle child visibility.")
    make_bullet(doc, "Indentation: Child rows are visually indented 24px per nesting level for clear hierarchy display.")
    doc.add_paragraph()

    make_heading(doc, "3.2  Spreadsheet-Style Inline Editing", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The matrix operates like a live spreadsheet — every cell is directly "
        "editable without opening a modal:", space_after=6)
    make_bullet(doc, "Double-click any cell to activate an inline input field (text, number, date, or dropdown).")
    make_bullet(doc, "Press Enter or click outside (blur) to save changes. Changes are immediately reflected in the dashboard.")
    make_bullet(doc, "Dropdown cells enforce valid system options — no free-text entry for controlled fields like Status, Priority, Complexity, or Module.")
    make_bullet(doc, "Ctrl+Z undoes the last inline edit within the session.")
    doc.add_paragraph()

    make_heading(doc, "3.3  Task Status Lifecycle", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Status transitions follow a controlled workflow. Not all transitions are "
        "permitted — the system enforces the following state machine:", space_after=6)

    add_data_table(doc,
        ["From Status", "Permitted Transitions", "Comment Required?"],
        [
            ["Not Started", "In Progress, Cancelled", "No"],
            ["In Progress", "Under Review, On Hold, Cancelled", "Required for On Hold, Cancelled"],
            ["Under Review", "Completed, In Progress, Cancelled", "Required for Cancelled"],
            ["On Hold", "In Progress, Under Review, Cancelled", "Required for Cancelled"],
            ["Completed", "In Progress (re-open)", "No"],
            ["Cancelled", "Not Started (re-activate)", "No"],
        ],
        col_widths=[1.6, 2.8, 2.0],
    )

    make_heading(doc, "3.4  Column Management & Visibility", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The visible column set, order, and column widths are fully configurable:", space_after=6)
    make_bullet(doc, "Right-click any column header to show the Column Visibility panel.")
    make_bullet(doc, "Drag column headers horizontally to reorder them.")
    make_bullet(doc, "Drag the right edge of a column header to resize it. All configurations persist across sessions via localStorage.")
    doc.add_paragraph()

    make_heading(doc, "3.5  Complexity & Effort Estimation", level=2, color=CLR_ACCENT)
    make_body(doc,
        "ProjectPulse uses a three-tier complexity model that scales estimated effort:", space_after=6)

    add_data_table(doc,
        ["Complexity Level", "Effort Multiplier", "Recommended Use Case"],
        [
            ["Easy",    "0.5×", "Well-understood tasks with clear requirements and low dependency."],
            ["Medium",  "1.0×", "Standard development tasks with moderate ambiguity."],
            ["Complex", "1.5×", "Tasks with significant unknowns, cross-team dependencies, or novel engineering challenges."],
        ],
        col_widths=[1.8, 1.8, 3.0],
    )

    add_callout(doc,
        "The Effort Unit (hrs / days / months) is a global system setting configured in Administration → General Settings. "
        "All effort values across the entire application honour this setting — switching units automatically recalculates all displayed values.",
        style="note")

    make_heading(doc, "3.6  Filtering & Sorting", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The left sidebar provides a powerful multi-dimensional filter panel:", space_after=6)
    make_bullet(doc, "Filter by Module, Status, Priority, Assignee, Release, Category, and Module Type simultaneously.")
    make_bullet(doc, "Click any column header once to sort ascending; click again to sort descending.")
    make_bullet(doc, "Active filters are displayed as removable pills at the top of the filter sidebar.")
    make_bullet(doc, "The search bar performs live full-text search across Task IDs, names, assignees, and notes.")
    doc.add_paragraph()

    add_page_break(doc)


def ch_gantt(doc):
    make_heading(doc, "4  Gantt Timeline & Dependency Management", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The Gantt Timeline provides a time-anchored visual representation of the "
        "project schedule. Tasks appear as horizontal bars spanning their start-to-due "
        "date range, with dependency arrows linking predecessor and successor tasks. "
        "Toggle to the Timeline view by clicking 'Timeline' in the Delivery Matrix "
        "toolbar.", space_after=8)

    add_screenshot(doc, "03_gantt_timeline",
                   "Figure 4.1 — Gantt Timeline with dependency arrows, critical path highlighting, and baseline comparison overlay")

    make_heading(doc, "4.1  Timeline Controls", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The timeline toolbar provides granular navigation controls:", space_after=6)
    make_bullet(doc, "Zoom Levels: Day / Week / Month / Quarter — adjust the time-axis resolution.")
    make_bullet(doc, "Scroll: Click-drag the timeline canvas to pan left/right across the schedule.")
    make_bullet(doc, "Today Line: A vertical amber marker indicates today's date on the timeline.")
    make_bullet(doc, "Baseline Overlay: Toggle the baseline date range (shown as a ghost bar behind each task bar) to visually compare planned vs. actual schedule positions.")
    doc.add_paragraph()

    make_heading(doc, "4.2  Dependency Arrows", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Finish-to-Start dependency relationships between tasks are rendered as "
        "curved SVG arrows connecting predecessor task bars to successor task bars.", space_after=6)
    make_bullet(doc, "Blue arrows: Active dependencies where the predecessor is still incomplete.")
    make_bullet(doc, "Green arrows: Resolved dependencies (predecessor is Completed).")
    make_bullet(doc, "Red arrows: Violated dependencies — the successor's start date precedes the predecessor's due date (scheduling conflict).")
    doc.add_paragraph()

    make_heading(doc, "4.3  Drag-to-Reschedule", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Task bars on the timeline are interactive:", space_after=6)
    make_bullet(doc, "Drag the bar body horizontally to shift the start and due date together (preserving duration).")
    make_bullet(doc, "Drag the right edge of the bar to extend or shorten the duration.")
    make_bullet(doc, "Cascading: When a task's dates shift, all downstream dependent tasks can optionally cascade their dates forward automatically.")
    doc.add_paragraph()

    add_callout(doc,
        "Cascading date shifts in the Gantt view operate on the live schedule by default. "
        "To experiment safely without affecting the live plan, activate Sandbox Mode first "
        "(available via the Scheduler view).",
        style="warning")

    add_page_break(doc)


def ch_scheduler(doc):
    make_heading(doc, "5  Intelligent Weekly Scheduler & Conflict Resolver", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The Weekly Scheduler is the resource-levelling engine of ProjectPulse. "
        "It maps team capacity against assigned task efforts across a rolling "
        "12-week horizon, surfacing over-allocations, dependency violations, and "
        "scheduling conflicts — and offering automated resolution heuristics via "
        "the Copilot engine.", space_after=8)

    add_screenshot(doc, "04_weekly_scheduler",
                   "Figure 5.1 — Weekly Scheduler: Resource Heatmap Grid, Conflict Diagnostics, and Copilot Cockpit")

    make_heading(doc, "5.1  The Weekly Heatmap Grid", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The central element of the Scheduler is a two-dimensional heatmap grid: "
        "team members on the vertical axis and calendar weeks on the horizontal axis. "
        "Each cell displays the resource's allocated hours vs. their weekly capacity:", space_after=6)

    add_data_table(doc,
        ["Cell State", "Display", "Condition", "Action Required"],
        [
            ["Over-allocated", "Neon crimson glow — e.g. '48h / 40h'", "Allocated hours > Capacity", "Reassign or extend task dates."],
            ["Optimal",        "Green highlight — e.g. '38h / 40h'",   "80%–100% of capacity used",  "No action required."],
            ["Under-utilised", "Slate grey — e.g. '22h / 40h'",        "< 70% of capacity used",      "Consider additional task assignments."],
            ["On Leave",       "Hatched pattern — '0h / 0h'",          "Leave recorded in Capacity Hub", "No tasks should be assigned."],
        ],
        col_widths=[1.4, 2.0, 1.8, 1.6],
    )

    make_heading(doc, "5.2  Left Navigation Sidebar", level=2, color=CLR_ACCENT)
    make_body(doc, "The scheduler sidebar provides four real-time widgets:", space_after=6)
    make_bullet(doc, "Statistics Bar: Shows Active Conflicts count, Monitored Tasks, Sandbox State (CLEAN / DIRTY), and current Week Offset.")
    make_bullet(doc, "Team Workload Widget: A progress-bar checklist showing total weekly allocated days per active team member.")
    make_bullet(doc, "Active Conflicts Widget: Stacks all detected scheduling overlaps and dependency violations with Lucide indicator icons.")
    make_bullet(doc, "Quick Tips: Contextual guidance on heatmap interpretation, sandbox usage, and date-shifting shortcuts.")
    doc.add_paragraph()

    make_heading(doc, "5.3  Right-Side Diagnostics Cockpit", level=2, color=CLR_ACCENT)
    make_body(doc, "A three-tab cockpit panel on the right side provides planning intelligence:", space_after=6)
    make_bullet(doc, "Diagnostics Tab: Lists every active conflict with its type (Over-allocated / Dependency Violation / Gap), affected resource, and impacted weeks.")
    make_bullet(doc, "Copilot Tab: Offers one-click automated resolution actions. Each action shows a preview of changes before applying.")
    make_bullet(doc, "Tips Tab: Displays keyboard shortcuts and sandbox state reminders.")
    doc.add_paragraph()

    make_heading(doc, "5.4  Automated Conflict Resolution Heuristics", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The autoResolveAllSchedulerConflicts() engine evaluates tasks in the "
        "sandbox state and applies three sequential heuristics:", space_after=6)

    add_data_table(doc,
        ["Heuristic", "Name", "Logic", "Pre-condition"],
        [
            ["A", "Auto-Sequence",
             "If Task_A → Task_B dependency exists and they overlap, Task_B's start date is shifted to Task_A's end date + 1 day.",
             "Both tasks must be in sandbox mode."],
            ["B", "Smart Reassignment",
             "If Resource R1 is over-allocated and Resource R2 shares the same Role with available capacity, tasks are re-assigned to R2.",
             "A role-matched resource with spare capacity must exist."],
            ["C", "Cascading Date Shift",
             "If a critical-path task is delayed, all transitively dependent future tasks are pushed forward, maintaining established buffer periods.",
             "Task dependency graph must be acyclic."],
        ],
        col_widths=[0.7, 1.6, 3.0, 1.3],
    )

    add_callout(doc,
        "All automated resolutions are applied to the Sandbox state only. "
        "Review the proposed changes in the Diagnostics tab before clicking 'Commit' "
        "to apply them to the live schedule.",
        style="warning")

    make_heading(doc, "5.5  Interactive Tooltips", level=2, color=CLR_ACCENT)
    make_body(doc, "Hovering over elements in the heatmap reveals live contextual data:", space_after=6)
    make_bullet(doc, "Grid Cells: Lists all tasks assigned to that resource for that week, total allocated hours, and remaining available hours.")
    make_bullet(doc, "Task Chips: Shows task status, parent deliverable name, priority level, and precise start/end dates.")
    doc.add_paragraph()

    add_page_break(doc)


def ch_raid(doc):
    make_heading(doc, "6  Unified RAID Register", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The RAID Register is the risk governance engine of ProjectPulse. It "
        "consolidates all Risks, Assumptions, Issues, and Dependencies into a "
        "single searchable register with formal exposure scoring, owner assignment, "
        "mitigation planning, and escalation tracking.", space_after=8)

    add_screenshot(doc, "05_raid_register",
                   "Figure 6.1 — Unified RAID Register with 5×5 Risk Matrix Heatmap and categorised RAID items")

    make_heading(doc, "6.1  RAID Classification Model", level=2, color=CLR_ACCENT)
    make_body(doc, "Every RAID entry is categorised into one of four types:", space_after=6)

    add_data_table(doc,
        ["Type", "Definition", "Primary Management Action"],
        [
            ["Risk",       "A potential future event that could negatively impact the project.",
             "Mitigate, transfer, or accept before the event occurs."],
            ["Assumption", "A factor believed to be true for planning purposes, without verified evidence.",
             "Validate through testing or stakeholder confirmation to eliminate associated risk."],
            ["Issue",      "An active problem currently impacting timeline, scope, or budget.",
             "Resolve immediately via assigned owner and mitigation strategy."],
            ["Dependency", "A reliance on an external team, system, third-party vendor, or deliverable.",
             "Map to task schedules. Monitor checkpoint dates. Escalate delays immediately."],
        ],
        col_widths=[1.3, 2.7, 2.4],
    )

    make_heading(doc, "6.2  Threat Exposure Scoring", level=2, color=CLR_ACCENT)
    make_body(doc,
        "For Risk-type items, threat exposure is calculated using the industry-standard "
        "5×5 probability-impact matrix:", space_after=6)
    make_body(doc, "    Exposure Score  =  Probability (1–5)  ×  Impact (1–5)", bold=True, space_after=8)

    add_data_table(doc,
        ["Score Range", "Rating", "Colour Indicator", "Management Protocol"],
        [
            ["15 – 25", "Critical",  "Red (🔴)",    "Immediate mitigation plan required. Executive escalation mandatory."],
            ["8 – 12",  "Medium",    "Amber (🟡)",  "Weekly monitoring. Mitigation plan to be prepared within 5 business days."],
            ["1 – 6",   "Low",       "Green (🟢)",  "Logged and reviewed bi-weekly. No immediate action required."],
        ],
        col_widths=[1.3, 1.2, 1.5, 3.4],
    )

    make_heading(doc, "6.3  RAID Item Lifecycle", level=2, color=CLR_ACCENT)
    make_body(doc, "RAID items progress through the following status workflow:", space_after=6)
    make_bullet(doc, "Identified: Newly logged item awaiting triage and owner assignment.")
    make_bullet(doc, "Active: Triaged and assigned — mitigation in progress.")
    make_bullet(doc, "Mitigated: Actions taken; item is under monitoring to confirm effectiveness.")
    make_bullet(doc, "Closed: Fully resolved with no residual risk.")
    make_bullet(doc, "Realized: Risk event has occurred and become an active Issue.")
    doc.add_paragraph()

    make_heading(doc, "6.4  Required Fields for Active RAID Items", level=2, color=CLR_ACCENT)
    make_body(doc, "To maintain governance standards, every Active RAID item must have all of the following:", space_after=6)
    make_numbered(doc, "Owner — An active team member assigned from the Capacity Hub.")
    make_numbered(doc, "Mitigation / Resolution Strategy — A detailed textual description of the action plan.")
    make_numbered(doc, "Target Closure Date — The deadline by which the issue must be resolved or the assumption validated.")
    doc.add_paragraph()

    add_callout(doc,
        "RAID items with Exposure Score ≥ 15 automatically contribute to the Health Index degradation "
        "on the Executive Overview Dashboard. Each unmitigated Critical risk deducts a fixed 5% from "
        "the Health Index.",
        style="caution")

    add_page_break(doc)


def ch_team(doc):
    make_heading(doc, "7  Team Capacity Hub & Role Management", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The Team Capacity Hub is the resource management control centre of "
        "ProjectPulse. It maintains team member profiles, capacity parameters, "
        "role assignments, utilisation statistics, and planned leave calendars. "
        "All scheduling and conflict detection in the Weekly Scheduler is derived "
        "from the capacity data managed here.", space_after=8)

    add_screenshot(doc, "06_team_capacity_hub",
                   "Figure 7.1 — Team Capacity Hub: Member profiles, utilisation bars, and leave calendar integration")

    make_heading(doc, "7.1  Adding and Managing Team Members", level=2, color=CLR_ACCENT)
    make_body(doc, "To add a new team member:", space_after=6)
    make_numbered(doc, "Navigate to the Team view via the top navigation bar.")
    make_numbered(doc, "Click '+ Add Member' in the top-right toolbar.")
    make_numbered(doc, "Enter the member's Name, Role, and initial Status.")
    make_numbered(doc, "Optionally configure Weekly Hour Cap (default: 40h) and Utilisation Rate (default: 100%).")
    make_numbered(doc, "Save. The member immediately appears in the Scheduler's resource grid.")
    doc.add_paragraph()

    make_heading(doc, "7.2  Capacity Parameters", level=2, color=CLR_ACCENT)
    make_body(doc, "Each team member has the following configurable capacity parameters:", space_after=6)

    add_data_table(doc,
        ["Parameter", "Default", "Description"],
        [
            ["Weekly Hour Cap",    "40h", "Maximum standard hours the member works per week."],
            ["Utilisation Rate",   "100%", "Percentage of standard time dedicated to project work. A member at 80% utilisation contributes 32 project hours per 40h week, reserving 8h for admin."],
            ["Status",             "Active", "Active members appear in Scheduler and can be assigned tasks. On Leave, Serving Notice, and Departed members are excluded from capacity calculations."],
            ["Planned Leaves",     "None", "Date ranges marked as leave. The Scheduler reduces that member's available hours to 0 for those weeks."],
        ],
        col_widths=[1.8, 1.2, 3.4],
    )

    make_heading(doc, "7.3  Dynamic Workload Balancing", level=2, color=CLR_ACCENT)
    make_body(doc, "The Capacity Hub displays a real-time utilisation indicator for each member:", space_after=6)
    make_bullet(doc, "Under-Utilised (< 70%): Blue indicator. The system suggests additional task assignments based on the member's role profile.")
    make_bullet(doc, "Balanced (70%–100%): Green indicator. Optimal resource loading.")
    make_bullet(doc, "Over-allocated (> 100%): Red warning flag. The Weekly Scheduler Diagnostics will list this resource as a conflict and recommend reassignment or date extension.")
    doc.add_paragraph()

    make_heading(doc, "7.4  Role-Based Assignment Logic", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Roles restrict which tasks a member can be automatically reassigned to during "
        "conflict resolution. The Smart Reassignment heuristic (Scheduler → Copilot) "
        "only swaps resources that share the same Role.", space_after=6)

    add_data_table(doc,
        ["Role", "Typical Responsibilities"],
        [
            ["Engineering Manager", "Project oversight, resource coordination, escalation management."],
            ["Lead Architect",      "System design, technical standards, code review ownership."],
            ["Lead Frontend",       "UI engineering, component library, performance optimisation."],
            ["Lead Design",         "UX research, design system, visual prototyping."],
            ["Developer",           "Feature implementation, unit testing, code review."],
            ["QA",                  "Test planning, defect logging, regression validation."],
            ["DevOps",              "Infrastructure, CI/CD pipelines, environment management."],
            ["Analyst",             "Requirements gathering, reporting, stakeholder communication."],
        ],
        col_widths=[2.2, 4.2],
    )

    add_page_break(doc)


def ch_defects(doc):
    make_heading(doc, "8  Defect Tracker & Bug Lifecycle", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The Defect Tracker enables QA teams and developers to log, triage, "
        "prioritise, assign, track, and close software defects through a "
        "formal lifecycle. Defect data is integrated into the Health Index "
        "calculation on the Executive Overview Dashboard and linked to "
        "specific tasks in the Delivery Matrix.", space_after=8)

    add_screenshot(doc, "07_defect_tracker",
                   "Figure 8.1 — Defect Tracker: Summary cards, severity-coded defect table, and bug lifecycle diagram")

    make_heading(doc, "8.1  Defect Severity Matrix", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Every defect must be assigned a severity level at the time of logging. "
        "Severity determines SLA response times and health index impact:", space_after=6)

    add_data_table(doc,
        ["Severity", "Code", "Definition", "Health Impact", "SLA Target"],
        [
            ["Blocker",  "S1", "Critical function completely broken. No workaround exists. Application unusable.",   "−10% per active Blocker", "24 hours"],
            ["Critical", "S2", "Major functionality impaired. Temporary workaround exists but is impractical.",       "−5% per active Critical",  "48 hours"],
            ["Major",    "S3", "Significant issue. Clear and stable workaround available.",                           "−2% per active Major",     "5 business days"],
            ["Minor",    "S4", "Trivial cosmetic issue, spelling error, or minor UI inconsistency.",                  "−0.5% per active Minor",   "10 business days"],
        ],
        col_widths=[1.1, 0.6, 2.4, 1.6, 1.4],
    )

    make_heading(doc, "8.2  Bug Lifecycle & Status Workflow", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Defects flow through the following mandatory status transitions. "
        "A defect cannot skip stages — the transition must be sequential:", space_after=6)

    add_data_table(doc,
        ["Status", "Description", "Responsible Party"],
        [
            ["New",      "Defect logged. Awaiting triage and owner assignment.",                                        "Reporter"],
            ["Assigned", "Owner allocated. Development investigation begun.",                                           "Engineering Manager"],
            ["Fixed",    "Code fix applied by development. Awaiting QA validation.",                                    "Developer"],
            ["Retest",   "Fix deployed to test environment. QA conducting verification testing.",                       "QA Engineer"],
            ["Closed",   "Fix verified by QA. Defect permanently closed.",                                              "QA Engineer"],
            ["Rejected", "Defect determined to be invalid (by design or not reproducible).",                           "Engineering Manager"],
            ["Deferred", "Fix postponed to a future release cycle with stakeholder approval.",                          "Project Manager"],
        ],
        col_widths=[1.3, 3.5, 1.6],
    )

    make_heading(doc, "8.3  Linking Defects to Tasks", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Every defect must be linked to a specific entity to maintain traceability:", space_after=6)
    make_bullet(doc, "Link Type: Task, Subtask, or GUI Screen.")
    make_bullet(doc, "Linked ID: The specific entity ID (e.g., TASK-004, ST-005).")
    make_bullet(doc, "The Delivery Matrix displays a defect count badge on the linked task row for instant visibility.")
    make_bullet(doc, "The Task Detail flyout lists all linked defects with their severity and status.")
    doc.add_paragraph()

    make_heading(doc, "8.4  Logging a New Defect — Step by Step", level=2, color=CLR_ACCENT)
    make_numbered(doc, "Navigate to the Defects view via the top navigation bar.")
    make_numbered(doc, "Click '+ Log Defect' in the top toolbar.")
    make_numbered(doc, "Complete the defect form: Title, Type, Severity, Priority, Linked Task, Assignee, Reporter, Description, and Reproduction Steps.")
    make_numbered(doc, "Click 'Save'. The defect is immediately added to the register with status 'New'.")
    make_numbered(doc, "The linked task's health contribution is recalculated and the dashboard updates in real time.")
    doc.add_paragraph()

    add_callout(doc,
        "S1 Blocker defects trigger an immediate red alert banner on the Executive Overview Dashboard. "
        "Blockers that remain open beyond their 24-hour SLA are escalated to the Engineering Manager "
        "via in-app notification.",
        style="caution")

    add_page_break(doc)


def ch_reports(doc):
    make_heading(doc, "9  Reports & Stakeholder Board Packs", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "ProjectPulse features a professional report generation engine that compiles "
        "real-time project data into presentation-ready board packs, structured Excel "
        "workbooks, and printable PDF status reports — all in one click.", space_after=8)

    add_screenshot(doc, "08_reports_boardpack",
                   "Figure 9.1 — Reports & Board Packs: Report builder, export controls, and preview panel")

    make_heading(doc, "9.1  Executive Board Pack", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The Board Pack compiles a complete project narrative suitable for executive "
        "and steering committee review:", space_after=6)
    make_bullet(doc, "Project Summary Narrative: Auto-generated high-level status paragraph, summarising key achievements, blockers, and next steps.")
    make_bullet(doc, "Aggregated Health Trend Graph: A Bezier-curve sparkline chart showing the Health Index over the past 8 weeks.")
    make_bullet(doc, "Critical Path Highlights: Tasks on the critical path that are within 5 days of or past their due dates.")
    make_bullet(doc, "Top 5 Risks: RAID items ranked by Exposure Score (Probability × Impact), with their owners and mitigation summaries.")
    make_bullet(doc, "Team Velocity Trend: 4-week rolling average of completed story points.")
    doc.add_paragraph()

    make_heading(doc, "9.2  Excel Workbook Export", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The full project workbook is generated using the ExcelJS library in the browser. "
        "The export is a multi-sheet workbook with professional formatting:", space_after=6)

    add_data_table(doc,
        ["Sheet", "Contents"],
        [
            ["System State",    "Project metadata, settings, effort unit, last saved timestamp, and file sync path."],
            ["Tasks",           "Full task register with all fields, baseline dates, variance calculations, and slippage reasons."],
            ["Team",            "Member directory with roles, statuses, capacity parameters, and leave records."],
            ["RAID Register",   "All Risk, Assumption, Issue, and Dependency items with exposure scores and mitigation details."],
            ["Defects",         "Complete defect log with severity, priority, status, linked tasks, and reproduction steps."],
            ["Activity Log",    "Last 500 audit log entries with timestamps, users, actions, and diff payloads."],
            ["Baselines",       "All captured baseline snapshots with per-task variance calculations."],
            ["Releases",        "Release version registry with status and target dates."],
        ],
        col_widths=[1.8, 4.6],
    )

    make_heading(doc, "9.3  Excel Cell Colour Coding", level=2, color=CLR_ACCENT)
    make_body(doc, "Exported Excel cells are colour-coded for instant readability:", space_after=6)

    add_data_table(doc,
        ["Colour", "Hex Code", "Applied To"],
        [
            ["Green",  "#D1E7DD / #0F5132", "Completed tasks, mitigated risks, closed defects, low RAID exposure."],
            ["Yellow", "#FFF3CD / #664D03", "In-progress tasks, medium risks, deferred defects."],
            ["Red",    "#F8D7DA / #842029", "Blocked / overdue tasks, critical risks, S1/S2 defects."],
            ["Grey",   "#F8F9FA / #6C757D", "Cancelled, Rejected, or On Hold items."],
        ],
        col_widths=[1.2, 2.0, 3.2],
    )

    add_callout(doc,
        "Exported spreadsheets contain live Excel formulas (SUM, AVERAGE, IF, COUNTIF) rather than "
        "static values where applicable. This means the workbook remains interactive and can be "
        "extended with additional analysis by stakeholders.",
        style="tip")

    add_page_break(doc)


def ch_activity(doc):
    make_heading(doc, "10  Audit Log & Activity Streams", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "ProjectPulse maintains an immutable, comprehensive audit trail of every "
        "state mutation — task CRUD operations, status transitions, defect lifecycle "
        "changes, RAID updates, and resource modifications. The Activity Log provides "
        "full accountability, change tracking, and session rollback capabilities.", space_after=8)

    add_screenshot(doc, "09_activity_audit_log",
                   "Figure 10.1 — Audit Log: Chronological activity stream with user identifiers and diff payloads")

    make_heading(doc, "10.1  What Gets Logged", level=2, color=CLR_ACCENT)
    make_body(doc, "Every significant state change generates an immutable log entry:", space_after=6)

    add_data_table(doc,
        ["Action Type", "Trigger", "Fields Captured"],
        [
            ["Created",        "New task, defect, RAID, or member added.",                   "Entity ID, name, all initial field values, timestamp, user."],
            ["Updated",        "Any field value changed on an existing entity.",              "Entity ID, field name, old value, new value, timestamp, user."],
            ["Status Changed", "Status field specifically transitions to a new value.",       "Entity ID, old status, new status, mandatory comment (for restricted transitions), timestamp, user."],
            ["Deleted",        "Task, defect, or RAID item permanently removed.",             "Entity ID, entity name, full field snapshot at time of deletion."],
        ],
        col_widths=[1.5, 2.5, 2.4],
    )

    make_heading(doc, "10.2  Log Entry Structure", level=2, color=CLR_ACCENT)
    make_body(doc, "Each log entry contains the following structured fields:", space_after=6)
    make_bullet(doc, "ts: ISO 8601 timestamp recorded to the millisecond.")
    make_bullet(doc, "user: Name of the active team member who performed the action.")
    make_bullet(doc, "action: Category of change — Created, Updated, Status Changed, or Deleted.")
    make_bullet(doc, "taskId / taskName: Reference to the primary entity affected.")
    make_bullet(doc, "field: The specific field that changed (e.g., 'status', 'dueDate', 'subtask.status').")
    make_bullet(doc, "oldVal / newVal: Structured diff payload capturing before and after values.")
    make_bullet(doc, "subtaskId: Optional — references a specific subtask if the change was at subtask level.")
    make_bullet(doc, "actCompletionDate: Optional — retrospective actual completion date for activity date resolution.")
    doc.add_paragraph()

    make_heading(doc, "10.3  Retrospective Activity Date Resolution", level=2, color=CLR_ACCENT)
    make_body(doc,
        "A key capability introduced in the June 2026 release is retrospective "
        "completion date resolution. Log entries may record an actCompletionDate "
        "that differs from the log insertion timestamp (ts). When this field is "
        "present, all dashboard widgets — including the 30-Day Activity Timeline, "
        "Weekly Velocity charts, Member Heatmaps, and Status Reports — use the "
        "actCompletionDate as the activity date rather than the log timestamp. "
        "This allows teams to accurately record when work was actually done, even "
        "if it was logged retrospectively days later.", space_after=8)

    make_heading(doc, "10.4  Session Rollback", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The Activity Log supports full session rollback to any previous state:", space_after=6)
    make_numbered(doc, "Locate the log entry representing the point you wish to restore.")
    make_numbered(doc, "Click 'Revert to this State' on the log entry row.")
    make_numbered(doc, "Confirm the reversion in the confirmation dialog.")
    make_numbered(doc, "ProjectPulse computes the inverse operations from the diff payloads of all subsequent entries and applies them sequentially.")
    make_numbered(doc, "The restored state is immediately persisted to localStorage via an automatic save() call.")
    doc.add_paragraph()

    add_callout(doc,
        "The Activity Log is auto-pruned to the last 500 entries to prevent unbounded localStorage growth. "
        "Export the log to Excel before reaching this limit if you need a permanent audit record beyond 500 entries.",
        style="warning")

    add_page_break(doc)


def ch_baselines(doc):
    make_heading(doc, "11  Schedule Baselines & Snapshot History", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "Schedule Baselines capture a static, approved snapshot of the project plan "
        "at a specific milestone — such as project kickoff, a phase gate, or after a "
        "major re-scope. Baselines serve as the immutable reference point for "
        "calculating schedule variance, analysing slippage, and restoring the plan "
        "if assumptions change.", space_after=8)

    add_screenshot(doc, "10_configuration_settings",
                   "Figure 11.1 — Schedule Baselines & Snapshot History panel in Project Settings")

    make_heading(doc, "11.1  Three Date Sets Explained", level=2, color=CLR_ACCENT)
    make_body(doc, "ProjectPulse tracks three parallel sets of dates for every task:", space_after=6)

    add_data_table(doc,
        ["Date Set", "Fields", "Purpose"],
        [
            ["Planned Dates",   "startDate, dueDate, estEffort",                  "The current working target dates that resources are actively working towards. These evolve during execution."],
            ["Baseline Dates",  "baselineStartDate, baselineDueDate, baselineEffort", "The frozen reference dates captured at a specific baseline snapshot point. Used to calculate variance. Immutable after capture."],
            ["Actual Dates",    "actCompletionDate, actEffort",                   "The real dates and effort recorded when tasks are completed. Used for EVM, velocity, and activity reporting."],
        ],
        col_widths=[1.5, 2.5, 2.4],
    )

    make_heading(doc, "11.2  Capturing a Baseline Snapshot", level=2, color=CLR_ACCENT)
    make_numbered(doc, "Open Project Settings by clicking the ⚙ gear icon in the top navigation bar.")
    make_numbered(doc, "Select 'Schedule Baselines' from the left sidebar within Settings.")
    make_numbered(doc, "Enter a Baseline Name (e.g., 'Initial Kickoff Plan', 'Sprint 3 Reforecast') and a Description.")
    make_numbered(doc, "Click 'Capture Snapshot'.")
    doc.add_paragraph()
    make_body(doc, "What happens when a snapshot is captured:", bold=True, space_after=4)
    make_bullet(doc, "All active tasks' baseline fields are overwritten with their current planned values.")
    make_bullet(doc, "A snapshot record is generated for every task and appended to the P.baselines array.")
    make_bullet(doc, "The snapshot is timestamped and named for future reference.")
    doc.add_paragraph()

    make_heading(doc, "11.3  Variance & Slippage Analysis", level=2, color=CLR_ACCENT)
    make_body(doc, "Once a baseline is captured, variance is automatically calculated and displayed:", space_after=6)

    add_data_table(doc,
        ["Metric", "Formula", "Interpretation"],
        [
            ["Variance Days",         "Planned/Actual Date − Baseline Date",       "+5d = 5 days behind schedule. −3d = 3 days ahead of schedule."],
            ["Effort Variance",       "estEffort − baselineEffort",                "Positive = scope increase. Negative = scope reduction."],
            ["Slippage Reason",       "Free-text field on the task flyout",        "Documented justification for schedule slippage. Captured in subsequent baseline snapshots and exported to Excel."],
        ],
        col_widths=[1.8, 2.2, 2.4],
    )

    make_heading(doc, "11.4  Restoring a Baseline", level=2, color=CLR_ACCENT)
    make_body(doc,
        "From the Schedule Baselines panel, select any historical snapshot and click "
        "'Restore'. This action:", space_after=6)
    make_bullet(doc, "Overwrites all current planned dates (startDate, dueDate, estEffort) with the snapshot's recorded values.")
    make_bullet(doc, "Resets baseline dates to match the restored snapshot.")
    make_bullet(doc, "Triggers a full recalculation of dashboard metrics, EVM figures, and schedule variance.")
    doc.add_paragraph()

    add_callout(doc,
        "Restoring a baseline is irreversible without another snapshot to return to. "
        "Always capture a 'Current State' baseline snapshot before restoring a historical one, "
        "so you can return to your pre-restore position if needed.",
        style="caution")

    add_page_break(doc)


def ch_config(doc):
    make_heading(doc, "12  System Configuration Reference", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The Administration panel (⚙ Settings) centralises all system-wide "
        "configuration options. Changes take effect immediately across all views "
        "without requiring a page reload.", space_after=8)

    add_screenshot(doc, "10_configuration_settings",
                   "Figure 12.1 — System Configuration: General Settings, Dropdown Manager, and Custom Fields panel")

    make_heading(doc, "12.1  General Settings", level=2, color=CLR_ACCENT)
    make_body(doc, "Core application behaviour is controlled by these settings:", space_after=6)

    add_data_table(doc,
        ["Setting", "Options", "Effect"],
        [
            ["Effort Unit",      "hrs / days / months", "Globally switches the unit label and rescales all effort values across the application."],
            ["Hours per Day",    "1–24 (default: 8)",   "Used to convert between hours and days when switching effort units."],
            ["Days per Week",    "1–7 (default: 5)",    "Defines the working week length for capacity calculations."],
            ["Working Days",     "Mon–Sun toggles",     "Specifies which days of the week are working days. Affects capacity calculations and date math."],
            ["Auto-Save",        "On / Off",            "Enables 10-second idle debounced auto-save to localStorage."],
        ],
        col_widths=[1.6, 2.0, 2.8],
    )

    make_heading(doc, "12.2  Dropdown Configuration Manager", level=2, color=CLR_ACCENT)
    make_body(doc,
        "All categorical dropdowns (Status, Priority, Category, Module, Role, etc.) "
        "are fully configurable. To modify a dropdown:", space_after=6)
    make_numbered(doc, "Open Settings → 'Dropdown Manager'.")
    make_numbered(doc, "Select the dropdown category to modify (e.g., 'Module').")
    make_numbered(doc, "Add new options, edit existing labels, or reorder items via drag-and-drop.")
    make_numbered(doc, "Click 'Save'. New options appear immediately in all views.")
    doc.add_paragraph()

    add_callout(doc,
        "Removing a dropdown option that is currently assigned to active tasks will not "
        "delete those tasks or their data — the old value is preserved on existing records. "
        "However, it will no longer appear as a selectable option for new assignments.",
        style="warning")

    make_heading(doc, "12.3  Custom Fields", level=2, color=CLR_ACCENT)
    make_body(doc,
        "Custom Fields allow teams to extend the Task and Defect schemas with "
        "project-specific attributes:", space_after=6)
    make_bullet(doc, "Supported Types: Text, Number, Date, Dropdown, Checkbox.")
    make_bullet(doc, "Scope: Custom fields can be defined independently for Tasks and for Defects.")
    make_bullet(doc, "Visibility: Custom fields appear as additional columns in the Delivery Matrix and Defect Tracker, and are included in Excel exports.")
    make_bullet(doc, "Persistence: Custom field values are stored alongside standard fields in localStorage and in the Excel workbook.")
    doc.add_paragraph()

    make_heading(doc, "12.4  Complexity Factor Multipliers", level=2, color=CLR_ACCENT)
    make_body(doc,
        "The default complexity multipliers (Easy: 0.5×, Medium: 1.0×, Complex: 1.5×) "
        "can be adjusted in Settings → 'Complexity Factors'. Changing these values "
        "recalculates all velocity and EVM metrics immediately.", space_after=6)

    make_heading(doc, "12.5  Theme & Visual Configuration", level=2, color=CLR_ACCENT)
    make_body(doc, "ProjectPulse ships with 20 curated visual themes:", space_after=6)
    make_bullet(doc, "Dark Themes: Nexus, Obsidian, Fintech, Codename, Workflow, Bento.")
    make_bullet(doc, "Light Themes: Emerald, Terracotta, Cloud, and 12 additional curated palettes.")
    make_bullet(doc, "Dark/Light Mode Override: Force dark or light mode regardless of the selected theme.")
    make_bullet(doc, "Theme selection is persisted per session and stored to the Excel workbook on export.")
    doc.add_paragraph()

    add_page_break(doc)


def ch_appendix(doc):
    make_heading(doc, "Appendix A — Keyboard Shortcuts", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc, "The following keyboard shortcuts are available globally across all views:", space_after=8)

    add_data_table(doc,
        ["Shortcut", "Action", "Scope"],
        [
            ["Ctrl + S",        "Force immediate save to localStorage and linked file.", "Global"],
            ["Ctrl + Z",        "Undo last inline edit.",                                "Delivery Matrix"],
            ["Ctrl + F",        "Focus the search / filter bar.",                        "Delivery Matrix, Defects"],
            ["Escape",          "Close open modal, flyout, or dropdown.",                "Global"],
            ["Enter",           "Confirm inline cell edit.",                             "Delivery Matrix"],
            ["Tab",             "Move to next editable cell.",                           "Delivery Matrix"],
            ["Shift + Tab",     "Move to previous editable cell.",                       "Delivery Matrix"],
            ["Arrow Keys",      "Navigate between cells in grid.",                       "Delivery Matrix"],
            ["Ctrl + Click",    "Multi-select tasks for bulk operations.",               "Delivery Matrix"],
            ["Space",           "Toggle expand/collapse on selected parent task.",       "Delivery Matrix"],
        ],
        col_widths=[1.5, 3.0, 1.9],
    )

    make_heading(doc, "Appendix B — Data Schema Reference", level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    make_body(doc,
        "The following tables document the full field-level schema for the primary "
        "data entities managed by ProjectPulse. All data is persisted in localStorage "
        "(pp-data key) and exported to the Excel workbook on save.", space_after=8)

    make_heading(doc, "B.1  Task Schema", level=2, color=CLR_ACCENT)
    add_data_table(doc,
        ["Field", "Type", "Description"],
        [
            ["id",                  "string",   "Unique identifier (e.g., TASK-001). Auto-generated."],
            ["name",                "string",   "Task title. Required."],
            ["status",              "string",   "Current status. Must match valid transition from VALID_TRANSITIONS."],
            ["priority",            "string",   "Critical / High / Medium / Low."],
            ["category",            "string",   "Classification from configured Category dropdown."],
            ["module",              "string",   "Associated module name from Module dropdown."],
            ["moduleType",          "string",   "Server / GUI / Interface."],
            ["assignee",            "string",   "Assigned member name."],
            ["release",             "string",   "Target release version (e.g., v1.0.0)."],
            ["startDate",           "YYYY-MM-DD", "Current planned start date."],
            ["dueDate",             "YYYY-MM-DD", "Current planned due date."],
            ["actCompletionDate",   "YYYY-MM-DD", "Actual completion date (null if incomplete)."],
            ["baselineStartDate",   "YYYY-MM-DD", "Frozen baseline start date."],
            ["baselineDueDate",     "YYYY-MM-DD", "Frozen baseline due date."],
            ["forecastDueDate",     "YYYY-MM-DD", "System-calculated forecast due date based on velocity."],
            ["slippageReason",      "string",   "Free-text explanation if task slips beyond baseline."],
            ["complexity",          "string",   "Easy / Medium / Complex."],
            ["progress",            "number",   "0–100 percentage completion."],
            ["estEffort",           "number",   "Estimated effort in active effortUnit."],
            ["actEffort",           "number",   "Actual logged effort in active effortUnit."],
            ["baselineEffort",      "number",   "Baseline effort in active effortUnit."],
            ["dependsOn",           "string[]", "Array of prerequisite Task IDs."],
            ["notes",               "string",   "Multi-line task description and notes."],
            ["subtasks",            "Subtask[]","Array of child step/screen items."],
            ["guiScreens",          "string[]", "Linked GUI screen names."],
            ["updatedAt",           "ISO 8601", "Timestamp of last modification."],
            ["parentId",            "string",   "Parent task ID for sub-rows (empty string for root tasks)."],
        ],
        col_widths=[1.8, 1.2, 3.4],
    )

    make_heading(doc, "B.2  Defect Schema", level=2, color=CLR_ACCENT)
    add_data_table(doc,
        ["Field", "Type", "Description"],
        [
            ["id",         "string",  "Unique identifier (e.g., DEF-001). Auto-generated."],
            ["title",      "string",  "Defect title. Required."],
            ["type",       "string",  "Functional Bug / UI/UX Issue / Performance / Security / Data Issue / Suggestion."],
            ["severity",   "string",  "S1 Blocker / S2 High / S3 Medium / S4 Low."],
            ["priority",   "string",  "Critical / High / Medium / Low."],
            ["status",     "string",  "New / Assigned / Fixed / Retest / Closed / Rejected / Deferred."],
            ["linkedType", "string",  "task / subtask / screen — type of linked entity."],
            ["linkedId",   "string",  "ID of the linked task, subtask, or screen."],
            ["assignee",   "string",  "Developer assigned to fix."],
            ["reporter",   "string",  "Person who logged the defect."],
            ["desc",       "string",  "Full defect description."],
            ["steps",      "string",  "Step-by-step reproduction instructions."],
        ],
        col_widths=[1.5, 1.0, 3.9],
    )


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN BUILD ORCHESTRATOR
# ══════════════════════════════════════════════════════════════════════════════

def build_docx():
    print(f"\n{'═'*60}")
    print("  ProjectPulse — Product Capabilities & User Manual")
    print(f"  Screenshots dir : {SS_DIR}")
    print(f"  Output          : {OUT_PATH}")
    print(f"{'═'*60}\n")

    doc = Document()
    set_page_margins(doc)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = CLR_TEXT

    # Heading styles
    for level, size in [(1, 18), (2, 13.5), (3, 11.5)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = CLR_NAVY
        h.font.bold = True

    print("  Building cover page…")
    build_cover(doc)

    print("  Building table of contents…")
    build_toc(doc)

    chapters = [
        ("Chapter 1 — Executive Overview Dashboard", ch_overview),
        ("Chapter 2 — Live Insights & Analytics",    ch_insights),
        ("Chapter 3 — Hierarchical Delivery Matrix", ch_delivery),
        ("Chapter 4 — Gantt Timeline",               ch_gantt),
        ("Chapter 5 — Weekly Scheduler",             ch_scheduler),
        ("Chapter 6 — RAID Register",                ch_raid),
        ("Chapter 7 — Team Capacity Hub",            ch_team),
        ("Chapter 8 — Defect Tracker",               ch_defects),
        ("Chapter 9 — Reports & Board Packs",        ch_reports),
        ("Chapter 10 — Audit Log",                   ch_activity),
        ("Chapter 11 — Schedule Baselines",          ch_baselines),
        ("Chapter 12 — Configuration Reference",     ch_config),
        ("Appendices",                               ch_appendix),
    ]

    for name, fn in chapters:
        print(f"  Writing {name}…")
        fn(doc)

    print(f"\n  Saving document → {OUT_PATH}")
    doc.save(OUT_PATH)
    print(f"\n✅  Done!  Document saved:\n   {OUT_PATH}\n")
    return OUT_PATH


if __name__ == "__main__":
    build_docx()
