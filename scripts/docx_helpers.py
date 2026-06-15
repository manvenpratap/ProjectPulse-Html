"""
docx_helpers.py
===============
Shared styling primitives, brand palette, and layout helpers used by all
three ProjectPulse document builders (build_pcd / build_fsd / build_umi).
"""

import os, glob, datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────
BASE_DIR  = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
SS_DIR    = os.path.join(BASE_DIR, "docs", "screenshots")
DOCS_DIR  = os.path.join(BASE_DIR, "docs")

# ── Brand Palette ──────────────────────────────────────────────────────────
CLR_NAVY    = RGBColor(0x0D, 0x1B, 0x2A)
CLR_TEAL    = RGBColor(0x00, 0xC2, 0xA8)
CLR_ACCENT  = RGBColor(0x1E, 0x9E, 0xD5)
CLR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
CLR_LIGHT   = RGBColor(0xF4, 0xF7, 0xFA)
CLR_DARK    = RGBColor(0x1A, 0x1A, 0x2E)
CLR_WARN    = RGBColor(0xE6, 0x7E, 0x22)
CLR_DANGER  = RGBColor(0xC0, 0x39, 0x2B)
CLR_SUCCESS = RGBColor(0x27, 0xAE, 0x60)
CLR_TEXT    = RGBColor(0x1C, 0x1C, 0x1E)
CLR_MUTED   = RGBColor(0x6C, 0x75, 0x7D)
CLR_BORDER  = RGBColor(0xD1, 0xD5, 0xDB)
CLR_PURPLE  = RGBColor(0x6C, 0x3F, 0xC2)
CLR_GOLD    = RGBColor(0xD4, 0xAF, 0x37)

# ── Document Setup ─────────────────────────────────────────────────────────

def new_document(top=1.0, bottom=1.0, left=1.25, right=1.25):
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)
    
    # Configure default normal style
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10.5)
    style.font.color.rgb = CLR_TEXT
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(8)
    
    # Configure Heading 1, 2, 3 styles globally
    heading_configs = [
        (1, 22, 24, 8),
        (2, 15, 16, 6),
        (3, 12, 12, 4)
    ]
    for level, size, before, after in heading_configs:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Segoe UI"
        h.font.size = Pt(size)
        h.font.color.rgb = CLR_NAVY
        h.font.bold = True
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.paragraph_format.keep_with_next = True
        
    return doc

# ── Cell / Table Helpers ───────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing_shd = tcPr.find(qn("w:shd"))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Sets inner margins (padding) of a table cell in twentieths of a point (dxa)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing_tcMar = tcPr.find(qn("w:tcMar"))
    if existing_tcMar is not None:
        tcPr.remove(existing_tcMar)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_borders(cell, color=None, sz=None, top="none", bottom="none", left="none", right="none", 
                     top_color="D1D5DB", bottom_color="D1D5DB", left_color="D1D5DB", right_color="D1D5DB",
                     top_sz="4", bottom_sz="4", left_sz="4", right_sz="4"):
    """
    Sets individual cell borders. Supports backwards compatibility for full borders
    when 'color' positional/keyword argument is passed.
    """
    if color is not None:
        top = bottom = left = right = "single"
        top_color = bottom_color = left_color = right_color = color
        if sz is not None:
            top_sz = bottom_sz = left_sz = right_sz = sz

    tcPr = cell._tc.get_or_add_tcPr()
    existing_borders = tcPr.find(qn("w:tcBorders"))
    if existing_borders is not None:
        tcPr.remove(existing_borders)
    tcBorders = OxmlElement("w:tcBorders")
    
    borders = [
        ("top", top, top_color, top_sz),
        ("bottom", bottom, bottom_color, bottom_sz),
        ("left", left, left_color, left_sz),
        ("right", right, right_color, right_sz)
    ]
    for edge, val, col, size in borders:
        tag = OxmlElement(f"w:{edge}")
        if val == "none" or val is None:
            tag.set(qn("w:val"), "none")
        else:
            tag.set(qn("w:val"), val)
            tag.set(qn("w:sz"), str(size))
            tag.set(qn("w:color"), col)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc, color="D1D5DB"):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_thick_rule(doc, color="00C2A8", sz="24"):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    return p


def add_page_break(doc):
    doc.add_page_break()

# ── Typography ─────────────────────────────────────────────────────────────

def make_heading(doc, text, level=1, color=None, space_before=None, space_after=None):
    h   = doc.add_heading(text, level=level)
    if space_before is None:
        space_before = 24 if level == 1 else (16 if level == 2 else 12)
    if space_after is None:
        space_after = 8 if level == 1 else (6 if level == 2 else 4)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.name = "Segoe UI"
    if color is None:
        color = CLR_NAVY
    run.font.color.rgb = color
    return h


def make_body(doc, text, space_before=0, space_after=6, color=None,
              bold=False, italic=False, size=10.5):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(size)
    if color:  run.font.color.rgb = color
    if bold:   run.bold  = True
    if italic: run.italic = True
    return p


def make_bullet(doc, text, level=0, space_after=3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(10.5)
    return p


def make_numbered(doc, text, level=0, space_after=3):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(10.5)
    return p


def make_formula(doc, text, space_before=6, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Inches(0.4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Courier New"
    run.font.bold = True
    run.font.color.rgb = CLR_ACCENT
    return p


def make_code(doc, text, space_before=4, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "Courier New"
    run.font.color.rgb = CLR_DARK
    return p

# ── Callout Boxes ──────────────────────────────────────────────────────────

def add_callout(doc, text, style="note"):
    configs = {
        "note":        ("INFO",         "0D4F8B", "E8F4FD", "1565C0"),
        "tip":         ("TIP",          "145A32", "E9F7EF", "1E8449"),
        "warning":     ("WARNING",      "7D4E00", "FEF9E7", "C67C00"),
        "caution":     ("CAUTION",      "6E2222", "FDEDEC", "C0392B"),
        "rule":        ("BUSINESS RULE","1A237E", "EDE7F6", "5E35B1"),
        "constraint":  ("CONSTRAINT",   "1B5E20", "E8F5E9", "2E7D32"),
        "integration": ("INTEGRATION",  "01579B", "E1F5FE", "0277BD"),
    }
    icons = {
        "note":        "ℹ️ ",
        "tip":         "💡 ",
        "warning":     "⚠️ ",
        "caution":     "🛑 ",
        "rule":        "📋 ",
        "constraint":  "🔒 ",
        "integration": "🔌 ",
    }
    label, label_col, bg_col, border_col = configs.get(style, configs["note"])
    icon_prefix = icons.get(style, "ℹ️ ")

    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style     = "Normal Table"
    
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_col)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    set_cell_borders(cell, left="single", left_color=border_col, left_sz="28")

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    lbl = p.add_run(f"{icon_prefix}{label}  ")
    lbl.bold = True
    lbl.font.size = Pt(9.5)
    lbl.font.name = "Segoe UI"
    lbl.font.color.rgb = RGBColor.from_string(label_col)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.line_spacing = 1.15
    body = p2.add_run(text)
    body.font.size = Pt(9.5)
    body.font.name = "Segoe UI"
    body.font.color.rgb = CLR_TEXT
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

# ── Data Tables ────────────────────────────────────────────────────────────

def add_data_table(doc, headers, rows, col_widths=None,
                   header_bg="0D1B2A", header_fg="FFFFFF"):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style     = "Normal Table"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=180, right=180)
        # Bold teal bottom border under header, none on others
        set_cell_borders(hdr_cells[i], bottom="single", bottom_color="00C2A8", bottom_sz="12")
        
        p   = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = "Segoe UI"
        run.font.color.rgb = RGBColor.from_string(header_fg)
        
    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx+1].cells
        bg    = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            set_cell_bg(cells[c_idx], bg)
            set_cell_margins(cells[c_idx], top=100, bottom=100, left=150, right=150)
            # Subtle horizontal dividers between rows, no vertical borders
            set_cell_borders(cells[c_idx], bottom="single", bottom_color="E2E8F0", bottom_sz="4")
            
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if p.runs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = "Segoe UI"
                    run.font.color.rgb = CLR_TEXT
                    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return tbl


def add_transition_table(doc, rows):
    return add_data_table(
        doc,
        ["From Status","To Status","Guard / Condition","Required Effect"],
        rows, col_widths=[1.5,1.5,2.2,1.4], header_bg="1A1A2E")


def add_schema_table(doc, rows):
    return add_data_table(
        doc,
        ["Field","Type","Constraint","Default","Description"],
        rows, col_widths=[1.4,0.8,1.1,0.8,2.5], header_bg="0D1B2A")

# ── Screenshot Embedding ───────────────────────────────────────────────────

def add_screenshot(doc, filename, caption, width_inches=6.3):
    candidates = [os.path.join(SS_DIR, filename),
                  os.path.join(SS_DIR, filename+".png")]
    path = None
    for c in candidates:
        if os.path.exists(c): path = c; break
    if not path:
        matches = glob.glob(os.path.join(SS_DIR, f"{filename}*"))
        if matches: path = sorted(matches)[0]

    is_diagram = "diagrams/" in filename or (path and "/diagrams/" in path)

    if path and os.path.exists(path):
        if not is_diagram:
            # Wrap regular UI screenshots in a clean frame table
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            set_cell_bg(cell, "FFFFFF")
            set_cell_borders(cell, color="E2E8F0", sz="4")
            
            p_img = cell.paragraphs[0]
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after  = Pt(6)
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(path, width=Inches(width_inches - 0.2))
        else:
            p_img = doc.add_paragraph()
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after  = Pt(0)
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(path, width=Inches(width_inches))
    else:
        p_img = doc.add_paragraph(f"[Screenshot: {filename}]")
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.runs[0].font.color.rgb = CLR_MUTED
        p_img.runs[0].font.italic    = True

    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(14)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = CLR_MUTED
    run.font.name = "Segoe UI"
    return p_img

# ── Cover Page ─────────────────────────────────────────────────────────────

def build_cover(doc, product_name, subtitle, doc_type, version,
                audience, confidentiality):
    doc.add_paragraph("\n\n")
    
    # Accent top border line
    add_thick_rule(doc, color="00C2A8", sz="36") # 4.5pt solid Teal bar
    
    # Left-aligned product name
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(18)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run(product_name)
    run.bold = True
    run.font.size = Pt(38)
    run.font.color.rgb = CLR_NAVY
    run.font.name = "Segoe UI"

    # Left-aligned subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_p.paragraph_format.space_before = Pt(4)
    sub_p.paragraph_format.space_after = Pt(28)
    run = sub_p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = CLR_TEAL
    run.font.name = "Segoe UI"

    doc.add_paragraph("\n")

    # Metadata Card block (single table cell with a thick left teal border and soft background)
    meta = [
        ("Document Type",   doc_type),
        ("Version",         version),
        ("Audience",        audience),
        ("Date",            datetime.date.today().strftime("%d %B %Y")),
        ("Confidentiality", confidentiality)
    ]
    
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Normal Table"
    
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "F4F7FA")
    set_cell_margins(cell, top=180, bottom=180, left=240, right=240)
    set_cell_borders(cell, left="single", left_color="00C2A8", left_sz="24") # 3pt left Teal border
    
    for idx, (label, value) in enumerate(meta):
        if idx == 0:
            p_meta = cell.paragraphs[0]
        else:
            p_meta = cell.add_paragraph()
        p_meta.paragraph_format.space_before = Pt(2)
        p_meta.paragraph_format.space_after = Pt(2)
        p_meta.paragraph_format.line_spacing = 1.15
        
        lbl_run = p_meta.add_run(f"{label}:  ")
        lbl_run.bold = True
        lbl_run.font.size = Pt(10)
        lbl_run.font.name = "Segoe UI"
        lbl_run.font.color.rgb = CLR_NAVY
        
        val_run = p_meta.add_run(value)
        val_run.font.size = Pt(10)
        val_run.font.name = "Segoe UI"
        val_run.font.color.rgb = CLR_TEXT
        
    add_page_break(doc)

# ── Table of Contents ──────────────────────────────────────────────────────

def build_toc(doc, chapters, title="Table of Contents"):
    make_heading(doc, title, level=1, color=CLR_NAVY)
    add_horizontal_rule(doc, "00C2A8")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    tbl = doc.add_table(rows=len(chapters), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Normal Table"
    
    for i, (num, title_text, pg) in enumerate(chapters):
        c0, c1, c2 = tbl.rows[i].cells
        c0.width = Inches(0.6)
        c1.width = Inches(5.0)
        c2.width = Inches(0.6)
        
        bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for cell in (c0, c1, c2):
            set_cell_bg(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            set_cell_borders(cell) # entirely borderless
            
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(num)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = CLR_TEAL
        r0.font.name = "Segoe UI"
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(title_text)
        r1.font.size = Pt(9.5)
        r1.font.name = "Segoe UI"
        r1.font.color.rgb = CLR_TEXT
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(pg)
        r2.bold = True
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = CLR_NAVY
        r2.font.name = "Segoe UI"
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_page_break(doc)

# ── Section Divider ────────────────────────────────────────────────────────

def add_section_divider(doc, part_label, part_title, description=""):
    add_page_break(doc)
    doc.add_paragraph("\n\n\n\n")
    
    # Left-aligned part label
    lbl_p = doc.add_paragraph()
    lbl_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lbl_p.paragraph_format.space_before = Pt(12)
    lbl_p.paragraph_format.space_after = Pt(4)
    run = lbl_p.add_run(part_label.upper())
    run.font.size = Pt(12)
    run.font.color.rgb = CLR_TEAL
    run.font.bold = True
    run.font.name = "Segoe UI"

    # Left-aligned large part title
    t_p = doc.add_paragraph()
    t_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_p.paragraph_format.space_before = Pt(4)
    t_p.paragraph_format.space_after = Pt(12)
    run = t_p.add_run(part_title)
    run.font.size = Pt(28)
    run.font.color.rgb = CLR_NAVY
    run.font.bold = True
    run.font.name = "Segoe UI"

    # Thick rule separator
    add_thick_rule(doc, color="00C2A8", sz="18") # 2.25pt

    if description:
        doc.add_paragraph("\n")
        d_p = doc.add_paragraph()
        d_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        d_p.paragraph_format.space_before = Pt(6)
        d_p.paragraph_format.space_after = Pt(12)
        run = d_p.add_run(description)
        run.font.size = Pt(11)
        run.font.color.rgb = CLR_MUTED
        run.font.italic = True
        run.font.name = "Segoe UI"
        
    add_page_break(doc)

# ── Document Control (FSD) ─────────────────────────────────────────────────

def add_document_control(doc, version="v2.1.0", status="Approved",
                          authors="ProjectPulse Product Team",
                          reviewers="Engineering Lead / QA Lead / Product Manager"):
    make_heading(doc,"Document Control",level=1,color=CLR_NAVY)
    add_horizontal_rule(doc,"00C2A8")
    add_data_table(doc,["Field","Value"],[
        ["Document Title",  "ProjectPulse - Functional Specification Document"],
        ["Version",         version],
        ["Status",          status],
        ["Author(s)",       authors],
        ["Reviewer(s)",     reviewers],
        ["Review Date",     datetime.date.today().strftime("%d %B %Y")],
        ["Classification",  "Internal - Restricted"],
        ["Next Review",     "Quarterly or upon major feature release"],
    ], col_widths=[2.0,4.4])
    add_page_break(doc)

# ── Footer ─────────────────────────────────────────────────────────────────

def add_footer(doc, doc_title):
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.paragraph_format.space_before = Pt(6)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"ProjectPulse  |  {doc_title}  |  Page ")
    run.font.size = Pt(8); run.font.color.rgb = CLR_MUTED
    run.font.name = "Segoe UI"
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"),"begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"),"end")
    run2 = para.add_run()
    run2.font.size=Pt(8); run2.font.color.rgb=CLR_MUTED
    run2.font.name = "Segoe UI"
    run2._r.append(fldChar1); run2._r.append(instrText); run2._r.append(fldChar2)
